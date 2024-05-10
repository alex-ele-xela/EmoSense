import io
import os

import warnings
warnings.filterwarnings("ignore")

import pickle
import subprocess

from transformers import RobertaTokenizer, RobertaConfig, RobertaForSequenceClassification
from spacy.lang.en import English
import docx
from docx.enum.text import WD_COLOR_INDEX

import cv2
from ultralytics import YOLO
from supervision import Detections
from torchvision import models

import torch
import torch.nn as nn
import librosa
from models.SpeechEmotionRecog import SpeechEmotionRecog

import numpy as np
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
sns.set_style("darkgrid")


import plotly.express as px
from PIL import Image, ImageDraw, ImageFont
from moviepy.editor import VideoFileClip, AudioFileClip

from tqdm import tqdm

import accelerate
from transformers import AutoProcessor, pipeline


def convert_video_to_audio_ffmpeg(video_path, audio_path):
    subprocess.call(["ffmpeg", "-y", "-i", video_path, audio_path],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.STDOUT)
    
    print(f"Successfully extracted audio file from video. \nAudio saved at {audio_path}\n")
    

class ASR():
    def __init__(self):
        print("Initializing Automatic Speech Recognition")

        asr_model_id = "openai/whisper-large-v3"

        processor = AutoProcessor.from_pretrained(asr_model_id)
        self.asr_model = pipeline(
            "automatic-speech-recognition",
            model=asr_model_id,
            tokenizer=processor.tokenizer,
            feature_extractor=processor.feature_extractor,
            max_new_tokens=128,
            chunk_length_s=30,
            batch_size=16,
            return_timestamps=True,
            device_map='auto',
        )
    
    def __call__(self, audio_path, output_dir):
        print("Processing audio file...")

        asr_result = self.asr_model(audio_path, generate_kwargs={"language": "english"})

        output_file_path = os.path.join(output_dir, "asr_result.pickle")
        with open(output_file_path, 'wb') as f:
            pickle.dump(asr_result, f)

        print(f"Successfully performed ASR. \nResult saved at {output_file_path}\n")

        return output_file_path


class TextEmotion():
    def __init__(self):
        print("Initializing Text Emotion Analysis")

        self.device = 'cuda' if torch.cuda.is_available() else 'cpu'
        with open('textemo_labeller.pickle', 'rb') as f:
            self.textemo_labeller = pickle.load(f)

        self.max_length = 256

        roberta_config = RobertaConfig(
            num_labels = len(self.textemo_labeller.classes_),
            output_attentions = False,
            output_hidden_states = False,
            max_position_embeddings = 514,
            type_vocab_size = 1
        )

        self.tokenizer = RobertaTokenizer.from_pretrained("roberta-base")

        self.sentiment_model = RobertaForSequenceClassification(config=roberta_config)
        self.sentiment_model.load_state_dict( torch.load('./models/sentiment_model.pth') )
        self.sentiment_model = self.sentiment_model.to(self.device)

        self.sentiment_to_color_mapper = {
            "angry": WD_COLOR_INDEX.RED,
            "disgust": WD_COLOR_INDEX.GREEN,
            "fear": WD_COLOR_INDEX.BLUE,
            "happy": WD_COLOR_INDEX.YELLOW,
            "neutral": WD_COLOR_INDEX.WHITE,
            "sad": WD_COLOR_INDEX.TEAL,
            "surprise": WD_COLOR_INDEX.VIOLET,
            "love": WD_COLOR_INDEX.TURQUOISE
        }

    def preprocess_text(self, text):
        return self.tokenizer(text, 
                              max_length = self.max_length, # 256 here
                              truncation=True, # truncate sequences longer than max_length to max_length
                              padding='max_length') # Adding padding to achieve max_length

    def get_pred(self, logits):
        softmax = nn.Softmax(dim=0)
        probs = softmax(logits)

        pred = torch.argmax(probs, dim=0)

        return pred.cpu().item()

    def get_sentiment_pred(self, text) -> list:
        encoding = self.preprocess_text(text)
        input_ids = torch.tensor(encoding['input_ids']).type(torch.long)
        attention_mask = encoding = torch.tensor(encoding['attention_mask']).type(torch.float)

        self.sentiment_model.eval()
        with torch.inference_mode():
            logits = self.sentiment_model(
                input_ids=input_ids.unsqueeze(0).to(self.device),
                attention_mask=attention_mask.unsqueeze(0).to(self.device)
            ).logits[0]

        sentiment_pred = self.get_pred(logits)
        sentiment_pred = self.textemo_labeller.inverse_transform([sentiment_pred])[0]

        return sentiment_pred
    
    def get_sent_list(self, text:str) -> list:
        nlp = English()
        nlp.add_pipe("sentencizer")
        doc = nlp(text.strip())

        sent_list = [str(sent) for sent in doc.sents]
        return sent_list

    def get_labelled_sents(self, sent_list:list) -> list:
        labelled_sents = []
        for sent in sent_list:
            sentiment = self.get_sentiment_pred(sent)

            labelled_sents.append((sent, sentiment))

        return labelled_sents
    
    def generate_document(self, labelled_sents:list) -> docx.document.Document:
        doc = docx.Document()

        doc.add_heading('Sentiment Labelled Transcript', 0)
        doc.add_paragraph("Please note that the sentiment highlighted here is solely calculated only on the basis of the text")

        doc.add_heading('Legend', 1)
        para = doc.add_paragraph()
        for sentiment, color in self.sentiment_to_color_mapper.items():
            para.add_run(color.name).font.highlight_color = color
            para.add_run(f" = {sentiment}\n")

        doc.add_heading('Transcript', 1)
        para = doc.add_paragraph()
        for sent, sentiment in labelled_sents:
            para.add_run(sent+" ").font.highlight_color = self.sentiment_to_color_mapper[sentiment]

        return doc

    def save_document(self, doc:docx.document.Document, file_path:str):
        doc.save(file_path)

    
    def __call__(self, asr_result_path, output_dir):
        print("Processing text...")

        with open(asr_result_path, 'rb') as f:
            self.asr_result = pickle.load(f)

        sent_list = self.get_sent_list(self.asr_result['text'])
        labelled_sents = self.get_labelled_sents(sent_list)

        doc = self.generate_document(labelled_sents)

        output_file_path = os.path.join(output_dir, "labelled_transcript.docx")
        self.save_document(doc, output_file_path)

        print(f"Successfully performed Text Emotion Analysis. \nLabelled Transcript saved at {output_file_path}\n")

        return output_file_path


class VideoTasks():
    def __init__(self, *args, **kwargs):
        pass

    def get_video(self, video_path):
        video = cv2.VideoCapture(video_path)

        if not video.isOpened():
            print("Error opening file")
            return

        fps = video.get(cv2.CAP_PROP_FPS)
        frame_count = int(video.get(cv2.CAP_PROP_FRAME_COUNT))

        return video, fps, frame_count

    def get_frame(self, video, frame_index):
        video.set(cv2.CAP_PROP_POS_FRAMES, frame_index)
        _, frame = video.read()

        rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        pil_frame = Image.fromarray(rgb_frame)

        return pil_frame


class FaceDetect(VideoTasks):
    def __init__(self, model_path, device):
        super(FaceDetect, self).__init__(self)

        print("Initializing Face Detection")

        self.device = device
        self.face_detect_model = YOLO(model_path, verbose=False).to(self.device)
    
    def get_bounds_confidence(self, frame):
        output = self.face_detect_model(frame)[0]
        results = Detections.from_ultralytics(output)

        return results.xyxy.tolist(), results.confidence
    
    def process_video(self, video_path):
        video, fps, frame_count = self.get_video(video_path)

        faces = {
            'frame_index': [],
            'timestamp': [],
            'x0': [],
            'y0': [],
            'x1': [],
            'y1': [],
            'confidence': []
        }
        for frame_index in range(frame_count):
            frame = self.get_frame(video, frame_index)
            bounds, confidence = self.get_bounds_confidence(frame)

            faces['frame_index'].append(frame_index)
            faces['timestamp'].append(frame_index/fps)
            if bounds:
                faces['x0'].append(bounds[0][0])
                faces['y0'].append(bounds[0][1])
                faces['x1'].append(bounds[0][2])
                faces['y1'].append(bounds[0][3])
                faces['confidence'].append(confidence[0])
            else:
                faces['x0'].append(0)
                faces['y0'].append(0)
                faces['x1'].append(0)
                faces['y1'].append(0)
                faces['confidence'].append(0)

        faces = pd.DataFrame(faces)

        return faces
    
    def __call__(self, video_path, output_dir):
        print("Processing video file...")

        faces = self.process_video(video_path)

        output_faces_path = os.path.join(output_dir, "faces.csv")
        faces.to_csv(output_faces_path, index=False)

        print(f"Successfully performed Face Detection. \nResult saved at {output_faces_path}\n")

        return output_faces_path


class FaceEmotion(VideoTasks):
    def __init__(self, model_path, device):
        print("Initializing Face Emotion Detection")

        self.device = device
        with open('face_emotion_labeller.pickle', 'rb') as f:
            self.face_emotion_labeller = pickle.load(f)

        self.face_emotion_model = models.vgg19()
        self.face_emotion_model.classifier[6] = nn.Linear(4096, len(self.face_emotion_labeller.classes_))

        self.face_emotion_model.load_state_dict(torch.load(model_path))
        self.face_emotion_model = self.face_emotion_model.to(self.device)

    def get_probs(self, logits):
        softmax = nn.Softmax(dim=1)
        probs = softmax(logits)

        return probs
    
    def get_facemo_scores(self, frame):
        logits = self.face_emotion_model(frame.to(self.device))
        probs = self.get_probs(logits).squeeze().cpu()

        scores = dict()
        for idx, prob in enumerate(probs):
            label = self.face_emotion_labeller.inverse_transform([idx])[0]
            scores[label]=prob.item()

        return scores
    
    def get_face_emotions(self, video_path, frame_indices, x0s, y0s, x1s, y1s):
        face_emotions = pd.DataFrame(
            {
                "frame_index": [],
                "angry": [],
                # "disgust": [],
                # "fear": [],
                "happy": [],
                "neutral": [],
                "sad": [],
                "surprise": []
            }
        )

        def get_facemo_row(scores, frame_index):
            new_row = {
                "frame_index": [frame_index, ],
                "angry": [scores.get("angry", 0), ],
                "disgust": [scores.get("disgust", 0), ],
                "fear": [scores.get("fear", 0), ],
                "happy": [scores.get("happy", 0), ],
                "neutral": [scores.get("neutral", 0), ],
                "sad": [scores.get("sad", 0), ],
                "surprise": [scores.get("surprise", 0), ]
            }

            return new_row

        video, _, _ = self.get_video(video_path)

        for frame_index, x0, y0, x1, y1 in zip(frame_indices, x0s, y0s, x1s, y1s):
            frame = self.get_frame(video, frame_index)
            if x0==0 and y0==0 and x1==0 and y1==0:
                face = frame
            else:
                face = frame.crop((x0, y0, x1, y1))

            face = face.resize((224,224))
            face = face.convert('RGB')

            face = np.asarray(face)
            h, w, c = face.shape
            face = torch.from_numpy(face).view(c,h,w).unsqueeze(0).type(torch.float)
            scores = self.get_facemo_scores(face)
            
            new_row = pd.DataFrame(get_facemo_row(scores, frame_index))
            face_emotions = pd.concat([face_emotions, new_row], axis=0)

        return face_emotions
    
    def plot_face_emotions(self, face_emotions):
        fig, axs = plt.subplots(3, 2, figsize=(10, 8))

        emotions = {
            "angry": "red",
            # "disgust": "green",
            # "fear": "blue",
            "happy": "yellow",
            "neutral": "black",
            "sad": "teal",
            "surprise": "purple",
            # "contempt": "magenta"
        }

        timestamps = face_emotions.frame_index * 22050
        i=0
        for emotion, color in emotions.items():
            sns.lineplot(
                x=timestamps,
                y=emotion,
                data=face_emotions,
                color=color,
                linewidth=1,
                ax=axs[i//2][i%2])

            axs[i//2][i%2].set_ylim(-0.1, 1.1)
            i+=1

        plt.suptitle("Facial Emotion Detection")

        return fig
    
    def __call__(self, video_path, faces_path, output_dir):
        print("Processing video file...")

        faces = pd.read_csv(faces_path)
        frame_indices = faces.frame_index.tolist()
        x0s = faces.x0.tolist()
        y0s = faces.y0.tolist()
        x1s = faces.x1.tolist()
        y1s = faces.y1.tolist()

        face_emotions = self.get_face_emotions(video_path, frame_indices, x0s, y0s, x1s, y1s)

        output_csv_path = os.path.join(output_dir, "face_emotions.csv")
        face_emotions.to_csv(output_csv_path, index=False)

        fig = self.plot_face_emotions(face_emotions)
        output_fig_path = os.path.join(output_dir, "face_emotion.png")
        fig.savefig(output_fig_path)

        print(f"Successfully performed Face Emotion Recognition. \nResult saved at {output_csv_path} and {output_fig_path}\n")

        return output_csv_path


class SER():
    def __init__(self, model_path, device):
        print("Initializing Speech Emotion Recognition")

        self.device = device

        with open('ser_emotion_labeller.pickle', 'rb') as f:
            self.ser_emotion_labeller = pickle.load(f)

        self.ser_model = SpeechEmotionRecog(num_labels=len(self.ser_emotion_labeller.classes_))
        self.ser_model.load_state_dict(torch.load(model_path))
        self.ser_model = self.ser_model.to(self.device)

        self.CLIP_LENGTH = 22050*5

    def set_sample_rate(audio, sample_rate, new_sample_rate):
        audio_resampled = librosa.resample(y=audio, 
                                        orig_sr=sample_rate, 
                                        target_sr=new_sample_rate)

        return audio_resampled

    def set_to_mono(self, audio):
        mono_audio = librosa.to_mono(audio)

        return mono_audio

    def preprocess_audio(self, audio, sample_rate):
        if sample_rate!=22050:
            audio= self.set_sample_rate(audio, sample_rate, 22050)

        audio_mono = self.set_to_mono(audio)

        return audio_mono

    def get_audio(self, audio_path):
        audio, sample_rate = librosa.load(audio_path)

        audio = self.preprocess_audio(audio, sample_rate)

        return audio
    
    def clip_audio(self, audio, start_idx, end_idx):
        clip = audio[start_idx:end_idx]

        return clip

    def get_probs(self, logits):
        softmax = nn.Softmax(dim=1)
        probs = softmax(logits)

        return probs

    def get_ser_scores(self, clip):
        audio = torch.tensor(clip).type(torch.float).view(1,1,self.CLIP_LENGTH)

        self.ser_model.eval()
        with torch.inference_mode():
            logits = self.ser_model(audio.to(self.device))

        probs = self.get_probs(logits).squeeze().cpu()

        scores = dict()
        for idx, prob in enumerate(probs):
            label = self.ser_emotion_labeller.inverse_transform([idx])[0]
            scores[label]=prob.item()

        return scores
    
    def get_ser_emotions(self, audio_path, timestamps):
        ser_emotions = pd.DataFrame(
            {
                "frame_index": [],
                "angry": [],
                "calm": [],
                "disgust": [],
                "fear": [],
                "happy": [],
                "neutral": [],
                "sad": [],
                "surprise": []
            }
        )

        def get_ser_row(scores, frame_index):
            new_row = {
                "frame_index": [frame_index, ],
                "angry": [scores.get("angry", 0), ],
                "calm": [scores.get("calm", 0), ],
                "disgust": [scores.get("disgust", 0), ],
                "fear": [scores.get("fear", 0), ],
                "happy": [scores.get("happy", 0), ],
                "neutral": [scores.get("neutral", 0), ],
                "sad": [scores.get("sad", 0), ],
                "surprise": [scores.get("surprise", 0), ]
            }

            return new_row

        audio = self.get_audio(audio_path)

        for frame_index, timestamp in enumerate(timestamps):
            end_idx = int(timestamp * 22050)
            start_idx = end_idx - self.CLIP_LENGTH

            if start_idx<0:
                scores = {}
            else:
                clip = self.clip_audio(audio, start_idx, end_idx)
                scores = self.get_ser_scores(clip)

            new_row = pd.DataFrame(get_ser_row(scores, frame_index))
            ser_emotions = pd.concat([ser_emotions, new_row], axis=0)

            if end_idx == len(audio):
                break

        return ser_emotions
    
    def plot_ser_emotions(self, ser_emotions):
        fig, axs = plt.subplots(3, 2, figsize=(10, 8))

        emotions = {
            "angry": "red",
            # "calm": "gray",
            "disgust": "green",
            "fear": "blue",
            "happy": "yellow",
            "neutral": "black",
            "sad": "teal",
            # "surprise": "purple"
        }

        timestamps = ser_emotions.frame_index * 22050
        i=0
        for emotion, color in emotions.items():
            sns.lineplot(
                x=timestamps,
                y=emotion,
                data=ser_emotions,
                color=color,
                linewidth=1,
                ax=axs[i//2][i%2])

            axs[i//2][i%2].set_ylim(-0.1, 1.1)
            i+=1

        plt.suptitle("Speech Emotion Recognition")

        return fig

    def save_fig(self, fig, figname):
        fig.savefig(figname)
    
    def __call__(self, audio_path, output_dir, timestamps):
        print("Processing audio...")
        ser_emotions = self.get_ser_emotions(audio_path, timestamps)

        output_csv_path = os.path.join(output_dir, "ser_emotions.csv")
        ser_emotions.to_csv(output_csv_path, index=False)

        fig = self.plot_ser_emotions(ser_emotions)
        output_fig_path = os.path.join(output_dir, "ser.png")
        fig.savefig(output_fig_path)

        print(f"Successfully performed SER. \nResults saved at {output_csv_path} and {output_fig_path}\n")

        return output_csv_path


class VideoMaker(VideoTasks):
    def __init__(self):
        print("Initializing Video Composition")
        super(VideoMaker, self).__init__(self)

    def get_drawn_frame(self, frame, bounds, confidence):
        draw_obj = ImageDraw.Draw(frame)
        font = ImageFont.truetype('arial.ttf', size=12)

        # Draw a rectangle with the bounding box coordinates
        draw_obj.rectangle(bounds, outline='blue', width=2)

        # Add a label with the confidence score
        text = f'face: {confidence:.1f}'
        text_size = draw_obj.textlength(text, font=font)
        text_x = bounds[2] - text_size
        text_y = bounds[1] - 12
        draw_obj.text((text_x, text_y), text, fill='red', font=font)

        return frame

    def get_video_frame(self, video, frame_index, row):
        frame = self.get_frame(video, frame_index)
        bounds = [row.x0.item(), row.y0.item(), row.x1.item(), row.y1.item()]
        confidence = row.confidence.item()
        if confidence>0:
            return self.get_drawn_frame(frame, bounds, confidence)
        else:
            return frame
    
    def get_face_emotions(self, path):
        df = pd.read_csv(path)
        face_emotions = {
            "emotion": [],
            "frame_index": [],
            "score": []
        }

        emotions = ['angry', 'happy', 'neutral', 'sad', 'surprise']

        for frame_index in df.frame_index.tolist():
            row = df[df['frame_index']==frame_index]
            for emotion in emotions:
                face_emotions['emotion'].append(emotion)
                face_emotions['frame_index'].append(frame_index)
                face_emotions['score'].append(row[emotion].item())
            

        face_emotions = pd.DataFrame(face_emotions)

        return face_emotions

    def get_ser_emotions(self, path):
        df = pd.read_csv(path)
        ser_emotions = {
            "emotion": [],
            "frame_index": [],
            "score": []
        }

        emotions = ['angry', 'disgust', 'fear', 'happy', 'neutral', 'sad']

        for frame_index in df.frame_index.tolist():
            row = df[df['frame_index']==frame_index]
            for emotion in emotions:
                ser_emotions['emotion'].append(emotion)
                ser_emotions['frame_index'].append(frame_index)
                ser_emotions['score'].append(row[emotion].item())
            

        ser_emotions = pd.DataFrame(ser_emotions)

        return ser_emotions
    
    def make_plot_function(self, width=None, height=None):
        if width and not height:
            width=width/100
            height=int(width * (6/10))
        elif height and not width:
            height=height/100
            width=int(height * (10/6))
        elif not width and not height:
            width=10
            height=6

        color_map = {
            "angry": "Red",
            "calm": "Pink", 
            "disgust": "Green",
            "fear": "Blue",
            "happy": "Yellow",
            "neutral": "White",
            "sad": "Teal",
            "surprise": "Violet"
        }

        def get_plot(df, title):
            fig = px.bar(df,
                     x='emotion',
                     y='score',
                     title=title,
                     color='emotion',
                     range_y=[0,1],
                     labels={
                         'emotion': 'Emotions',
                         'score': 'Scores'
                         },
                     color_discrete_map=color_map,
                     width=width,
                     height=height
                     )
            img_data = fig.to_image(format="png")
            fig = Image.open(io.BytesIO(img_data))

            return fig
        
        return get_plot
    
    def get_frame_face_ser(self, video, get_plot, frame_index, faces, face_emotions, ser_emotions):
        face_row = faces[faces['frame_index']==frame_index]
        video_frame = self.get_video_frame(video, frame_index, face_row)
        
        face = face_emotions[face_emotions['frame_index']==frame_index]
        face_fig = get_plot(face, "Face Emotion Recognition")

        ser = ser_emotions[ser_emotions['frame_index']==frame_index]
        ser_fig = get_plot(ser, "Speech Emotion Recognition")

        return video_frame, face_fig, ser_fig
    
    def get_composite_frame(self, video_frame, face_fig, ser_fig, width, height):
        composite = Image.new('RGB', (width, height), (200, 200, 200))

        video_x, video_y = video_frame.size
        video_frame_offset = (0, (height-video_y)//2)
        face_fig_offset = (width-face_fig.size[0], 0)
        ser_fig_offset = (width-ser_fig.size[0], height-ser_fig.size[1])

        composite.paste(video_frame, video_frame_offset)
        composite.paste(face_fig, face_fig_offset)
        composite.paste(ser_fig, ser_fig_offset)

        return composite
    
    def __call__(self, video_path, audio_path, output_dir, faces_path, face_emotions_path, ser_emotions_path):
        print("Compositing video...")

        faces = pd.read_csv(faces_path)
        face_emotions = self.get_face_emotions(face_emotions_path)
        ser_emotions = self.get_ser_emotions(ser_emotions_path)

        video, fps, _ = self.get_video(video_path)

        orig_width, orig_height = self.get_frame(video, 0).size
        frame_width = int(orig_width * (2/3))
        frame_height = int(frame_width * (orig_height/orig_width))

        plot_width = orig_width - frame_width - 2
        plot_height = (orig_height-2) // 2
        get_plot = self.make_plot_function(width=plot_width, height=plot_height)

        fourcc = cv2.VideoWriter_fourcc(*'mp4v')
        video_writer = cv2.VideoWriter('temp.mp4', fourcc, fps, (orig_width, orig_height))

        frame_indices = faces.frame_index.tolist()

        for frame_index in tqdm(frame_indices):
            video_frame, face_fig, ser_fig = self.get_frame_face_ser(video, get_plot, frame_index, faces, face_emotions, ser_emotions)
            video_frame = video_frame.resize((frame_width, frame_height))
            
            composite_frame = self.get_composite_frame(video_frame, face_fig, ser_fig, orig_width, orig_height)   
            composite_frame = np.array(composite_frame)[:,:,::-1]

            video_writer.write(composite_frame)

        video_writer.release()

        video_clip = VideoFileClip('temp.mp4')
        audio_clip = AudioFileClip(audio_path)
        video_clip = video_clip.set_audio(audio_clip)

        output_video_path = os.path.join(output_dir, 'result.mp4')
        video_clip.write_videofile(output_video_path)

        os.remove('temp.mp4')

        print(f"Successfully created video. \nResult saved at {output_video_path}\n")

        return output_video_path


def main(VIDEO_PATH): 
    OUTPUT_DIR = 'outputs\\' + VIDEO_PATH.split('.')[0]
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    AUDIO_PATH = os.path.join(OUTPUT_DIR, "audio.wav")

    device = 'cuda' if torch.cuda.is_available() else 'cpu'

    convert_video_to_audio_ffmpeg(VIDEO_PATH, AUDIO_PATH)

    asr = ASR()
    asr_result_path = asr(AUDIO_PATH, OUTPUT_DIR)

    text_emotion = TextEmotion()
    transcript_path = text_emotion(asr_result_path, OUTPUT_DIR)
    del text_emotion

    face_detect = FaceDetect(model_path='models/face_detect_model.pt',
                             device=device)
    faces_path = face_detect(VIDEO_PATH, OUTPUT_DIR)
    del face_detect

    face_emotion = FaceEmotion(model_path='models/face_emotion_model.pth',
                               device=device)
    face_emotions_path = face_emotion(VIDEO_PATH, faces_path, OUTPUT_DIR)
    del face_emotion

    timestamps = pd.read_csv(faces_path).timestamp.tolist()
    ser = SER(model_path='models/ser_model.pth',
              device=device)
    ser_emotions_path = ser(AUDIO_PATH, OUTPUT_DIR, timestamps)
    del ser

    video_maker = VideoMaker()
    output_video_path = video_maker(VIDEO_PATH, AUDIO_PATH, OUTPUT_DIR, faces_path, face_emotions_path, ser_emotions_path)


if __name__=="__main__":
    main('mahjong.mp4')